from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "updated_paper_full.md"
DOCX_PATH = BASE_DIR / "updated_paper_full.docx"


def set_global_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_inline_code(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_markdown_like_paragraph(doc: Document, line: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    chunks = re.split(r"(`[^`]+`)", line)
    for chunk in chunks:
        if chunk.startswith("`") and chunk.endswith("`"):
            add_inline_code(p, chunk[1:-1])
        else:
            p.add_run(chunk)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    if i == 0:
                        run.bold = True
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_image(doc: Document, alt_text: str, image_rel_path: str):
    image_path = (BASE_DIR / image_rel_path).resolve()
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(14.5))
    caption = doc.add_paragraph(alt_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_md_lines(md_text: str):
    return md_text.splitlines()


def build_docx():
    doc = Document()
    set_global_style(doc)
    lines = parse_md_lines(MD_PATH.read_text(encoding="utf-8"))

    table_buffer = []
    in_table = False

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            if in_table and table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if len(cells) > 1 and all(set(c.replace("-", "").replace(":", "").strip()) == set() for c in cells):
                continue
            in_table = True
            table_buffer.append(cells)
            continue
        else:
            if in_table and table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.add_run(stripped[2:])
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            p.add_run(re.sub(r"^\d+\.\s*", "", stripped))
            continue

        if stripped == "---":
            doc.add_paragraph("")
            continue

        # Keep math blocks as plain text to avoid rendering issues in docx.
        if stripped.startswith("\\[") or stripped.endswith("\\]") or stripped.startswith("\\text") or stripped.startswith("R("):
            add_text_paragraph(doc, stripped)
            continue

        add_markdown_like_paragraph(doc, line)

    if in_table and table_buffer:
        add_table(doc, table_buffer)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
